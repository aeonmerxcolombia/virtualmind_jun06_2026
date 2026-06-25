import os, json, requests, uuid, mimetypes
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query
from fastapi.responses import FileResponse, Response
from sqlalchemy.orm import Session
from typing import Optional, List
from datetime import datetime

from app.database.db import SessionLocal
from app.auth.jwt_handler import verify_token
from app.models.hoja_vida import HojaVida
from app.models.user import User
from app.models.competencia import Competencia
from app.models.project import Project
from app.schemas.hoja_vida_schema import HojaVidaOut, HojaVidaUpdate

router = APIRouter(prefix="/rrhh", tags=["RRHH"])

UPLOAD_DIR = "/home/ubuntu/backend/documentos/hojas_vida"
os.makedirs(UPLOAD_DIR, exist_ok=True)

from app.services.ai.gemini_pool import get_gemini_key

GEMINI_MODEL = "gemini-2.5-flash"


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def parsear_hoja_vida_con_gemini(texto: str) -> dict:
    if not texto.strip():
        return {"error": "No hay API key o texto vacío"}

    prompt = f"""Eres un asistente de RRHH experto en parsear hojas de vida y sugerir mejoras de perfil.
Analiza el siguiente texto extraído de una hoja de vida y devuelve SOLO un JSON válido con esta estructura exacta (sin texto adicional, sin markdown):

{{
  "nombre_completo": "",
  "email": "",
  "telefono": "",
  "direccion": "",
  "perfil_profesional": "",
  "habilidades": [],
  "experiencia": [{{"cargo": "", "empresa": "", "fecha_inicio": "", "fecha_fin": "", "descripcion": ""}}],
  "educacion": [{{"titulo": "", "institucion": "", "fecha": "", "nivel": ""}}],
  "idiomas": [],
  "certificaciones": [],
  "sugerencias_perfil": {{
    "cargo_sugerido": "",
    "empresa_sugerida": "",
    "interes_principal_sugerido": "",
    "formato_preferido_sugerido": "",
    "nivel_experiencia_sugerido": "",
    "objetivo_principal_sugerido": ""
  }}
}}

Completa TODOS los campos que puedas identificar.
Para sugerencias_perfil, deduce del perfil profesional y experiencia:
- cargo_sugerido: el cargo más reciente o relevante
- empresa_sugerida: la empresa más reciente
- interes_principal_sugerido: el área técnica de mayor afinidad (elige entre: desarrollo_software, diseno_instruccional, diseno_grafico, produccion_multimedia, gestion_proyectos, marketing, inteligencia_artificial, ciencia_datos, ciberseguridad, educacion, comunicacion, audiovisual, animacion_3d, investigacion, innovacion)
- formato_preferido_sugerido: según perfil (elige entre: video, lectura, interactivo, practico, auditivo, mixto)
- nivel_experiencia_sugerido: según años de experiencia (elige entre: principiante, intermedio, avanzado, senior)
- objetivo_principal_sugerido: según metas profesionales (elige entre: aprendizaje, actualizacion, certificacion, mejora_continua, nuevo_empleo, emprendimiento)

Si un campo no tiene información, déjalo vacío o como Array vacío.
NO inventes información.

Texto de la hoja de vida:
{texto[:15000]}"""

    try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={get_gemini_key()}"
        body = {"contents": [{"parts": [{"text": prompt}]}]}
        resp = requests.post(url, json=body, timeout=60)
        data = resp.json()
        text = (
            data.get("candidates", [{}])[0]
            .get("content", {})
            .get("parts", [{}])[0]
            .get("text", "")
            .strip()
        )
        text = text.replace("```json", "").replace("```", "").strip()
        return json.loads(text)
    except Exception as e:
        return {"error": str(e), "raw": text if "text" in dir() else ""}


def extraer_texto_archivo(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    text = ""
    try:
        if ext == ".pdf":
            import pdfplumber

            with pdfplumber.open(filepath) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            if len(text.strip()) < 50:
                try:
                    import pytesseract
                    from PIL import Image
                    import io
                    import pypdfium2 as pdfium

                    pdf_doc = pdfium.PdfDocument(filepath)
                    ocr_pages = []
                    for i in range(len(pdf_doc)):
                        page = pdf_doc[i]
                        bitmap = page.render(scale=3)
                        pil_image = bitmap.to_pil()
                        ocr_text = pytesseract.image_to_string(
                            pil_image, lang="spa+eng"
                        )
                        ocr_pages.append(ocr_text)
                    ocr_result = "\n".join(ocr_pages)
                    if len(ocr_result.strip()) > len(text.strip()):
                        text = ocr_result
                except Exception:
                    pass
        elif ext == ".docx":
            from docx import Document

            doc = Document(filepath)
            text = "\n".join(p.text for p in doc.paragraphs)
        elif ext == ".txt":
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            text = "Formato no soportado para extracción de texto"
    except ImportError as e:
        text = f"Error: Librería no disponible - {e}"
    except Exception as e:
        text = f"Error al extraer texto: {e}"
    return text


@router.post("/hoja-vida/upload")
async def subir_hoja_vida(
    file: UploadFile = File(...),
    project_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    token_data: dict = Depends(verify_token),
):
    user_id_token = token_data.get("user_id", "")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in (".pdf", ".docx", ".txt"):
        raise HTTPException(status_code=400, detail="Solo se aceptan PDF, DOCX o TXT")

    stored_name = f"{uuid.uuid4()}{ext}"
    filepath = os.path.join(UPLOAD_DIR, stored_name)
    content = await file.read()
    with open(filepath, "wb") as f:
        f.write(content)

    texto_extraido = extraer_texto_archivo(filepath)

    datos_parseados = parsear_hoja_vida_con_gemini(texto_extraido)

    if not datos_parseados or "error" in datos_parseados:
        hv = HojaVida(
            user_id=user_id_token,
            project_id=project_id,
            filename_original=file.filename,
            filename_almacenado=stored_name,
            ruta_archivo=filepath,
        )
        db.add(hv)
        db.commit()
        return {
            "mensaje": "Archivo subido pero no se pudo parsear automáticamente",
            "id": hv.id,
            "error_parse": datos_parseados.get("error", "Desconocido"),
        }

    email_extraido = datos_parseados.get("email", "").strip().lower()

    usuario_existente = None
    if email_extraido:
        usuario_existente = db.query(User).filter(User.email == email_extraido).first()

    hv = HojaVida(
        user_id=user_id_token,
        project_id=project_id,
        nombre_completo=datos_parseados.get("nombre_completo"),
        email=email_extraido,
        telefono=datos_parseados.get("telefono"),
        direccion=datos_parseados.get("direccion"),
        perfil_profesional=datos_parseados.get("perfil_profesional"),
        habilidades=datos_parseados.get("habilidades"),
        experiencia=datos_parseados.get("experiencia"),
        educacion=datos_parseados.get("educacion"),
        idiomas=datos_parseados.get("idiomas"),
        certificaciones=datos_parseados.get("certificaciones"),
        filename_original=file.filename,
        filename_almacenado=stored_name,
        ruta_archivo=filepath,
    )
    db.add(hv)
    db.commit()
    db.refresh(hv)

    respuesta = {
        "mensaje": "Hoja de vida procesada exitosamente",
        "id": hv.id,
        "datos_extraidos": {
            "nombre_completo": hv.nombre_completo,
            "email": hv.email,
            "telefono": hv.telefono,
            "habilidades": hv.habilidades,
            "experiencia": hv.experiencia,
            "educacion": hv.educacion,
        },
        "sugerencias_perfil": datos_parseados.get("sugerencias_perfil", {}),
    }

    if usuario_existente:
        respuesta["usuario_existente"] = {
            "uid": usuario_existente.uid,
            "nombre": usuario_existente.nombre,
            "email": usuario_existente.email,
            "mensaje": f"El usuario {usuario_existente.nombre} ya está registrado. ¿Deseas actualizar sus datos?",
        }

    return respuesta


@router.get("/hoja-vida/{hv_id}", response_model=HojaVidaOut)
def obtener_hoja_vida(
    hv_id: int,
    db: Session = Depends(get_db),
    token_data: dict = Depends(verify_token),
):
    hv = db.query(HojaVida).filter(HojaVida.id == hv_id).first()
    if not hv:
        raise HTTPException(status_code=404, detail="Hoja de vida no encontrada")
    return hv


@router.get("/hoja-vida/{hv_id}/view")
def ver_hoja_vida(
    hv_id: int,
    db: Session = Depends(get_db),
    token_data: dict = Depends(verify_token),
):
    hv = db.query(HojaVida).filter(HojaVida.id == hv_id).first()
    if not hv or not hv.ruta_archivo:
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    if not os.path.exists(hv.ruta_archivo):
        raise HTTPException(
            status_code=404, detail="Archivo físico no encontrado en el servidor"
        )
    media_type, _ = mimetypes.guess_type(hv.ruta_archivo)
    return FileResponse(
        hv.ruta_archivo,
        media_type=media_type or "application/octet-stream",
        filename=hv.filename_original,
    )


@router.get("/hoja-vida/{hv_id}/download")
def descargar_hoja_vida(
    hv_id: int,
    db: Session = Depends(get_db),
    token_data: dict = Depends(verify_token),
):
    hv = db.query(HojaVida).filter(HojaVida.id == hv_id).first()
    if not hv or not hv.ruta_archivo:
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    if not os.path.exists(hv.ruta_archivo):
        raise HTTPException(
            status_code=404, detail="Archivo físico no encontrado en el servidor"
        )
    return FileResponse(
        hv.ruta_archivo,
        media_type="application/octet-stream",
        filename=hv.filename_original,
    )


@router.get("/archivo/{hv_id}")
def archivo_hoja_vida(
    hv_id: int,
    db: Session = Depends(get_db),
    token_data: dict = Depends(verify_token),
):
    hv = db.query(HojaVida).filter(HojaVida.id == hv_id).first()
    if not hv or not hv.ruta_archivo:
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    if not os.path.exists(hv.ruta_archivo):
        raise HTTPException(
            status_code=404, detail="Archivo físico no encontrado en el servidor"
        )
    media_type, _ = mimetypes.guess_type(hv.ruta_archivo)
    return FileResponse(
        hv.ruta_archivo,
        media_type=media_type or "application/octet-stream",
        filename=hv.filename_original,
    )


@router.get("/hoja-vida/usuario/{user_id}")
def obtener_hv_por_usuario(
    user_id: str,
    db: Session = Depends(get_db),
    token_data: dict = Depends(verify_token),
):
    hv_list = (
        db.query(HojaVida)
        .filter(HojaVida.user_id == user_id)
        .order_by(HojaVida.fecha_subida.desc())
        .all()
    )
    return hv_list


@router.get("/hoja-vida/proyecto/{project_id}")
def listar_hv_por_proyecto(
    project_id: int,
    db: Session = Depends(get_db),
    token_data: dict = Depends(verify_token),
):
    roles = token_data.get("roles", [])
    if not any(r in roles for r in ("superadmin", "admin", "coordinador", "cliente")):
        raise HTTPException(status_code=403, detail="No autorizado")
    hv_list = (
        db.query(HojaVida)
        .filter(HojaVida.project_id == project_id)
        .order_by(HojaVida.fecha_subida.desc())
        .all()
    )
    return hv_list


@router.get("/hoja-vida/")
def listar_todas_hv(
    db: Session = Depends(get_db),
    token_data: dict = Depends(verify_token),
):
    roles = token_data.get("roles", [])
    if "superadmin" not in roles:
        raise HTTPException(status_code=403, detail="Solo superadmin")
    return db.query(HojaVida).order_by(HojaVida.fecha_subida.desc()).all()


@router.put("/hoja-vida/{hv_id}")
def actualizar_hoja_vida(
    hv_id: int,
    data: HojaVidaUpdate,
    db: Session = Depends(get_db),
    token_data: dict = Depends(verify_token),
):
    hv = db.query(HojaVida).filter(HojaVida.id == hv_id).first()
    if not hv:
        raise HTTPException(status_code=404, detail="No encontrada")

    update_data = data.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        setattr(hv, key, value)

    db.commit()
    return {"ok": True, "mensaje": "Hoja de vida actualizada"}


@router.put("/hoja-vida/{hv_id}/vincular-usuario")
def vincular_usuario(
    hv_id: int,
    user_id: str = Form(...),
    db: Session = Depends(get_db),
    token_data: dict = Depends(verify_token),
):
    hv = db.query(HojaVida).filter(HojaVida.id == hv_id).first()
    if not hv:
        raise HTTPException(status_code=404, detail="No encontrada")
    user = db.query(User).filter(User.uid == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Usuario no encontrado")

    hv.user_id = user_id
    db.commit()

    competencia = db.query(Competencia).filter(Competencia.user_id == user_id).first()
    if not competencia:
        competencia = Competencia(user_id=user_id)
        db.add(competencia)

    if hv.habilidades:
        competencia.habilidades = hv.habilidades
    if hv.educacion:
        niveles = {
            "pregrado": "Profesional",
            "especializacion": "Especialización",
            "maestria": "Maestría",
            "doctorado": "Doctorado",
        }
        niv = None
        for e in hv.educacion or []:
            if isinstance(e, dict) and e.get("nivel"):
                niv = niveles.get(e["nivel"].lower(), e["nivel"])
        if niv:
            competencia.nivel_academico = niv
        areas = []
        for e in hv.educacion or []:
            if isinstance(e, dict) and e.get("titulo"):
                areas.append(e["titulo"])
        if areas:
            competencia.area_conocimiento = "; ".join(areas)
    if hv.idiomas:
        competencia.idiomas = hv.idiomas
    if hv.perfil_profesional:
        competencia.perfilamiento = hv.perfil_profesional
    if hv.experiencia:
        total_anios = sum(
            _calcular_anios_exp(e)
            for e in (hv.experiencia or [])
            if isinstance(e, dict)
        )
        competencia.anios_experiencia = total_anios

    db.commit()
    return {
        "ok": True,
        "mensaje": "Hoja de vida vinculada al usuario y competencias actualizadas",
    }


def _calcular_anios_exp(exp: dict) -> int:
    try:
        inicio = exp.get("fecha_inicio", "")
        fin = exp.get("fecha_fin", "")
        if fin and inicio:
            year_diff = int(fin[:4]) - int(inicio[:4])
            return max(year_diff, 1)
        return 1
    except:
        return 1


@router.get("/buscar-usuario")
def buscar_usuario_por_email(
    email: str = Query(...),
    db: Session = Depends(get_db),
    token_data: dict = Depends(verify_token),
):
    user = db.query(User).filter(User.email == email.strip().lower()).first()
    if not user:
        return {"existe": False}
    return {
        "existe": True,
        "uid": user.uid,
        "nombre": user.nombre,
        "email": user.email,
    }


@router.get("/check-email")
def check_email_rrhh(
    email: str = Query(...),
    db: Session = Depends(get_db),
):
    email_lower = email.strip().lower()
    user = db.query(User).filter(User.email == email_lower).first()
    hv = (
        db.query(HojaVida)
        .filter(HojaVida.email == email_lower)
        .order_by(HojaVida.fecha_subida.desc())
        .first()
    )
    return {
        "usuario_existe": bool(user),
        "usuario": {"uid": user.uid, "nombre": user.nombre, "email": user.email}
        if user
        else None,
        "hv_existe": bool(hv),
        "hv": {
            "id": hv.id,
            "nombre_completo": hv.nombre_completo,
            "telefono": hv.telefono,
            "habilidades": hv.habilidades,
            "perfil_profesional": hv.perfil_profesional,
            "experiencia": hv.experiencia,
            "educacion": hv.educacion,
        }
        if hv
        else None,
    }
